from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION
from pathlib import Path
from re import compile, search, escape, match, IGNORECASE
from config import *
import traceback, logging
import time

logging.basicConfig(format='%(message)s', level=debug_level)

word_count = 0

def formatDocument(input, output):
    global word_count
    title_added = False
    file_name = file.name.replace(".docx", "").strip()

    document = Document(input)

    author_name = document.core_properties.author
    created = document.core_properties.created
    created_year = "" if type(created) == type(None) else created.year

    ## Clean document
    # Remove empty paragraphs
    def deleteParagraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Add subtitle after title
    def addSubtitle(paragraph):
        subtitle = document.add_paragraph(copyrightText(created_year, author_name), style='Subtitle')
        subt = subtitle._p
        p = paragraph._p
        p.addnext(subt)

    def capitalizeSentences(text):
        text_list = text.split()
        for i in range(len(text_list)):
            text_list[i] = text_list[i].lower()
            if (not text_list[i] in capitalize_words_list or i == 0):
                text_list[i] = text_list[i].capitalize()

        text = ' '.join(text_list)
        return text

    # Set style using runs
    def runSetStyle(paragraph, style):
        for run in paragraph.runs:
            italic, bold, underline = run.italic, run.bold, run.underline
            run.font.name = style.font.name
            run.font.size = style.font.size
            run.font.color.rgb = style.font.color.rgb
            run.italic, run.bold, run.underline = italic, bold, underline
        

    ### Format styles
    styles = document.styles

    ## Format title
    if ('Title' in document.styles):
        title_style = document.styles['Title']
    else:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)    

    title_style.hidden = False
    title_style.quick_style = True
    title_style.priority = 1

    title_style.paragraph_format.alignment = title_paragraph_alignment
    title_style.paragraph_format.page_break_before = title_paragraph_page_break_before
    title_style.paragraph_format.space_before = title_paragraph_space_before
    title_style.paragraph_format.space_after = title_paragraph_space_after
    title_style.font.name = title_font_name
    title_style.font.size = title_font_size
    title_style.font.color.rgb = title_font_color
    title_style.font.bold = title_font_bold
    title_style.font.italic = title_font_italic
    title_style.font.underline = title_font_underline

    ## Format chapters    
    if ('Heading 1' in document.styles):
        heading_style = document.styles['Heading 1']
    else:
        heading_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

    heading_style.hidden = False
    heading_style.quick_style = True
    heading_style.priority = 2

    heading_style.paragraph_format.alignment = heading_1_paragraph_alignment
    heading_style.paragraph_format.page_break_before = heading_1_paragraph_page_break_before
    heading_style.paragraph_format.space_before = heading_1_paragraph_space_before
    heading_style.paragraph_format.space_after = heading_1_paragraph_space_after
    heading_style.font.name = heading_1_font_name
    heading_style.font.size = heading_1_font_size
    heading_style.font.color.rgb = heading_1_font_color
    heading_style.font.bold = heading_1_font_bold
    heading_style.font.italic = heading_1_font_italic
    heading_style.font.underline = heading_1_font_underline

    ## Format normal
    if ('Normal' in document.styles):
        normal_style = document.styles['Normal']
    else:
        normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)

    normal_style.hidden = False
    normal_style.quick_style = True
    normal_style.priority = 0

    normal_style.paragraph_format.first_line_indent = normal_paragraph_first_line_indent
    normal_style.paragraph_format.left_indent = normal_paragraph_left_indent
    normal_style.paragraph_format.right_indent = normal_paragraph_right_indent
    normal_style.paragraph_format.alignment = normal_paragraph_alignment
    normal_style.paragraph_format.space_before = normal_paragraph_space_before
    normal_style.paragraph_format.space_after = normal_paragraph_space_after
    normal_style.paragraph_format.page_break_before = normal_paragraph_page_break_before
    normal_style.font.name = normal_font_name
    normal_style.font.size = normal_font_size
    normal_style.font.color.rgb = normal_font_color
    normal_style.font.bold = normal_font_bold
    normal_style.font.italic = normal_font_italic
    normal_style.font.underline = normal_font_underline

    ## Format subtitle
    if ('Subtitle' in document.styles):
        subtitle_style = document.styles['Subtitle']
    else:
        subtitle_style = styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    
    normal_style.hidden = False
    normal_style.quick_style = True
    normal_style.priority = 3

    subtitle_style.base_style = document.styles[subtitle_inherits_from]
    subtitle_style.paragraph_format.alignment = subtitle_paragraph_alignment
    subtitle_style.paragraph_format.space_after = subtitle_paragraph_space_after
    subtitle_style.font.name = subtitle_font_name
    subtitle_style.font.size = subtitle_font_size
    subtitle_style.font.bold = subtitle_font_bold
    subtitle_style.font.italic = subtitle_font_italic
    subtitle_style.font.underline = subtitle_font_underline
    
    for para in document.paragraphs:
        para_text = para.text.strip()
        para_text_old = para_text
        list_of_actions_logs = ""

        para.paragraph_format.first_line_indent = None
        para.paragraph_format.left_indent = None
        para.paragraph_format.right_indent = None

        if (para_text != ""):
            if (not title_added):
                if (para.style.name != title_style.name):
                    para_text = file_name
                    list_of_actions_logs = list_of_actions_logs + " [Title]"

                para.style = title_style
                para_text = capitalizeSentences(para_text)
                addSubtitle(para)
                title_added = True
                para.text = para_text

            # Some people use Title instead of Heading 1
            elif (para.style.name == title_style.name):
                para.style = heading_style
                para.text = para_text
                list_of_actions_logs = list_of_actions_logs + " [Style is Title]"

            # Check for Heading 1 text (starting with header_1_names_list or numeric value and max 75 characters)
            if ((len(para_text) <= CHAPTER_MAX_LENGTH) or (para.style.name == heading_style.name)):
                ## Find elements in paragraph text
                # List of header 1 keywords present at the beginning of the text (empty or one word only)
                header_1_keyword_first = [ele for ele in header_1_names_list if para_text.upper().startswith(ele)]
                
                if (header_1_keyword_first and len(para_text.split()) >= 2):
                    # List of letter numbers (whole word only with eventually . or : at the end) | (?i) = case insensitive search
                    letter_number = [ele for ele in number_dict.keys() if search(r"(?i)(?<!\S)" + escape(ele) + r"[\.:]{0,1}" + r"(?!\S)", para_text.split()[1])]   
                    # List of first digits in text (with . and : characters stuck to it)
                    digit = [ele for ele in para_text if match(r"(?<!\S)" + r"\d+" + r"[\.:]{0,1}" + r"(?!\S)", para_text.split()[1])]             
                else:
                    # List of letter numbers (whole word only with eventually . or : at the end) | (?i) = case insensitive search
                    letter_number = [ele for ele in number_dict.keys() if search(r"(?i)(?<!\S)" + escape(ele) + r"[\.:]{0,1}" + r"(?!\S)", para_text.split()[0])]
                    # List of first digits in text (with . and : characters stuck to it)
                    digit = [ele for ele in para_text if match(r"(?<!\S)" + r"\d+" + r"[\.:]{0,1}" + r"(?!\S)", para_text.split()[0])]
                    
                # If there is a chapter name and header_1_keyword and/or chapter number set to Heading 1
                if ((not header_1_keyword_first) and (letter_number or digit) and len(para_text.split()) > 1):
                    para.style = heading_style
                    para_text = para_text.replace('.', '')
                    para_text = para_text.replace(':', '')
                    para_text = header_1_names_list[0].capitalize() + " " + para_text
                    # header_1_keyword_first, digit, letter_number = [], [], []
                    list_of_actions_logs = list_of_actions_logs + " [No header keyword + number]"
                    para.text = para_text
                elif (header_1_keyword_first and (letter_number or digit) and len(para_text.split()) > 2):
                    para.style = heading_style
                    para_text = para_text.replace('.', '')
                    para_text = para_text.replace(':', '')
                    # header_1_keyword_first, digit, letter_number = [], [], []
                    list_of_actions_logs = list_of_actions_logs + " [Header keyword + number]"
                    para.text = para_text

                # If whole text is a number (digit)
                if (digit and para_text == digit):
                    para.style = heading_style
                    para.text = para_text = header_1_names_list[0].capitalize() + " " + para_text
                    list_of_actions_logs = list_of_actions_logs + " [Whole Text = Number]"

                # If whole text is a number (in letter) convert it to number
                elif (letter_number and para_text.upper() == letter_number[0]):
                    para.style = heading_style
                    para_text = header_1_names_list[0].capitalize() + " " + str(number_dict[letter_number[0]])
                    list_of_actions_logs = list_of_actions_logs + " [Text = Letter number]"
                    para.text = para_text

                # Replace chapter name number in letter with the corresponding number
                elif (header_1_keyword_first):
                    para.style = heading_style

                    if (letter_number):
                        for substring in number_dict.keys():
                            if substring in para_text.upper():
                                chapter_number_found = substring
                        if (chapter_number_found):
                            pattern = compile(chapter_number_found, IGNORECASE)
                            para_text = pattern.sub(str(number_dict[chapter_number_found.upper()]), para_text)
                    list_of_actions_logs = list_of_actions_logs + " [Letter to number]"
                    para.text = para_text

                # If no conditions were met, apply normal style
                if (para.style != heading_style and para.style != title_style and para.style != subtitle_style):
                    runSetStyle(para, normal_style)
                    para.style = normal_style
                else:
                    para.text = capitalizeSentences(para_text)

            else:
                runSetStyle(para, normal_style)
                para.style = normal_style

        else:
            deleteParagraph(para)

        if (para_text_old != para_text):
            logging.debug("%s \"%s\" --> \"%s\"", list_of_actions_logs, para_text_old, para_text)
        word_count = word_count + len(para_text.split())

    # Document sections iteration
    # Remove headers/footers and set correct page orientation/format
    for section in document.sections:
        if (section.header and not keep_headers):
            section.header.is_linked_to_previous = True
        if (section.footer and not keep_footers):
            section.footer.is_linked_to_previous = True
        if (section.orientation != page_orientation):
            logging.info("[%s] Switching page orientation", file_name)
            section.orientation = page_orientation
        section.page_height = page_height
        section.page_width = page_width

        section.top_margin = top_margin
        section.bottom_margin = bottom_margin
        section.left_margin = left_margin
        section.right_margin = right_margin

    # Save document
    document.save(output.format(word_count=word_count))

start_time = time.time()

# Find all docx files in input folder and recreate subfolders in output_folder
files_list = list(Path().glob(input_folder + "/**/*.docx"))
for file in files_list:
    input_file_path = file.as_posix()
    temp_output_file_path = f"{output_folder}/{file.relative_to(*file.parts[:1]).as_posix()}"
    if (file.parents[-2] != output_folder):
        logging.info("\nWorking on %s", input_file_path)
        Path(temp_output_file_path).parents[0].mkdir(parents=True, exist_ok=True)
        try:
            output_file_path = temp_output_file_path.replace(".docx", " - {word_count}.docx")
            formatDocument(input_file_path, output_file_path)
            if (word_count == 0):
                logging.warning("No words were detected in %s (document might be a table)\n", file.name.replace("docx", ""))
            word_count = 0
        except Exception:
            traceback.print_exc()
            logging.error("    /!\ %s failed /!\ \n", input_file_path)

logging.info("\n========== Finished in %ss ==========", (time.time() - start_time))
variable = input('Press enter to exit')